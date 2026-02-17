"""Writer-category fixture file generators (docx, odt, rtf, html, txt)."""

from io import BytesIO
from pathlib import Path
from typing import Optional


# ---------------------------------------------------------------------------
# Shared test data — "제품 사양서" scenario (FR-016)
# ---------------------------------------------------------------------------

WRITER_TITLE = "LibreFormer 제품 사양"

WRITER_PARAGRAPHS = [
    (
        "LibreFormer는 Python 기반의 문서 변환 라이브러리입니다. "
        "LibreOffice를 백엔드로 활용하여 다양한 문서 포맷 간 변환을 지원합니다. "
        "간단한 API를 통해 docx, xlsx, pptx, odt 등의 파일을 PDF로 변환할 수 있습니다."
    ),
    (
        "주요 기능으로는 단일 파일 변환, 배치 변환, 병렬 처리가 있습니다. "
        "변환 결과는 Succeed 또는 Failed 타입으로 반환되어 타입 안전한 에러 처리가 가능합니다. "
        "LibreOffice의 자동 설치 기능도 Linux 환경에서 지원됩니다."
    ),
    (
        "CI/CD 파이프라인에서도 안정적으로 동작하도록 설계되었습니다. "
        "세션 스코프 fixture를 통해 테스트 성능을 최적화하고, "
        "외부 파일 의존 없이 순수 Python 코드로 테스트 파일을 생성합니다."
    ),
]

WRITER_TABLE_HEADERS = ["기능", "설명", "상태"]

WRITER_TABLE_ROWS = [
    ["문서 변환", "단일 파일을 대상 포맷으로 변환", "✅ 완료"],
    ["배치 처리", "여러 파일을 동시에 변환", "✅ 완료"],
    ["비동기 지원", "asyncio 기반 비동기 변환 API", "✅ 완료"],
]


# ---------------------------------------------------------------------------
# T007: create_docx
# ---------------------------------------------------------------------------


def create_docx(path: Path, image_bytes: Optional[bytes] = None) -> Path:
    """Create a .docx file with heading, paragraphs, table, and optional image.

    Uses the "제품 사양서" (Product Specification) scenario data (FR-016).

    Args:
        path: File path to create.
        image_bytes: Optional PNG bytes to embed as inline image.

    Returns:
        The created file Path.
    """
    from docx import Document
    from docx.shared import Inches

    doc = Document()

    # Heading
    doc.add_heading(WRITER_TITLE, level=1)

    # Paragraphs
    for para in WRITER_PARAGRAPHS:
        doc.add_paragraph(para)

    # Table (3×3)
    doc.add_heading("사양 비교", level=2)
    table = doc.add_table(rows=1, cols=len(WRITER_TABLE_HEADERS), style="Table Grid")
    header_cells = table.rows[0].cells
    for i, header in enumerate(WRITER_TABLE_HEADERS):
        header_cells[i].text = header
    for row_data in WRITER_TABLE_ROWS:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = cell_text

    # Inline image
    if image_bytes is not None:
        doc.add_heading("제품 이미지", level=2)
        doc.add_picture(BytesIO(image_bytes), width=Inches(2))

    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# T008: create_odt
# ---------------------------------------------------------------------------


def create_odt(path: Path) -> Path:
    """Create a .odt file with heading, paragraphs, and table.

    Requires odfpy. Caller should handle ImportError via pytest.importorskip.

    Args:
        path: File path to create.

    Returns:
        The created file Path.
    """
    from odf.opendocument import OpenDocumentText
    from odf.text import H, P
    from odf.table import Table, TableRow, TableCell

    doc = OpenDocumentText()

    # Heading
    h = H(outlinelevel=1, text=WRITER_TITLE)
    doc.text.addElement(h)

    # Paragraphs
    for para_text in WRITER_PARAGRAPHS:
        p = P(text=para_text)
        doc.text.addElement(p)

    # Table
    h2 = H(outlinelevel=2, text="사양 비교")
    doc.text.addElement(h2)

    table = Table(name="SpecTable")
    # Header row
    header_row = TableRow()
    for header in WRITER_TABLE_HEADERS:
        cell = TableCell()
        cell.addElement(P(text=header))
        header_row.addElement(cell)
    table.addElement(header_row)
    # Data rows
    for row_data in WRITER_TABLE_ROWS:
        row = TableRow()
        for cell_text in row_data:
            cell = TableCell()
            cell.addElement(P(text=cell_text))
            row.addElement(cell)
        table.addElement(row)
    doc.text.addElement(table)

    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# T009: create_rtf
# ---------------------------------------------------------------------------


def create_rtf(path: Path) -> Path:
    """Create a .rtf file with RTF markup. No external library needed.

    Args:
        path: File path to create.

    Returns:
        The created file Path.
    """
    lines = [r"{\rtf1\ansi\deff0"]
    # Title
    lines.append(r"{\b\fs36 " + WRITER_TITLE + r"}\par\par")
    # Paragraphs
    for para in WRITER_PARAGRAPHS:
        lines.append(para + r"\par\par")
    # Table header
    lines.append(r"{\b " + r"\tab ".join(WRITER_TABLE_HEADERS) + r"}\par")
    # Table rows
    for row in WRITER_TABLE_ROWS:
        lines.append(r"\tab ".join(row) + r"\par")
    lines.append("}")

    path.write_text("\n".join(lines), encoding="utf-8")
    return path


# ---------------------------------------------------------------------------
# T010: create_html
# ---------------------------------------------------------------------------


def create_html(path: Path) -> Path:
    """Create a .html file with HTML markup. No external library needed.

    Args:
        path: File path to create.

    Returns:
        The created file Path.
    """
    paragraphs_html = "\n".join(f"<p>{p}</p>" for p in WRITER_PARAGRAPHS)
    headers_html = "".join(f"<th>{h}</th>" for h in WRITER_TABLE_HEADERS)
    rows_html = "\n".join(
        "<tr>" + "".join(f"<td>{c}</td>" for c in row) + "</tr>"
        for row in WRITER_TABLE_ROWS
    )

    html = f"""<!DOCTYPE html>
<html lang="ko">
<head><meta charset="utf-8"><title>{WRITER_TITLE}</title></head>
<body>
<h1>{WRITER_TITLE}</h1>
{paragraphs_html}
<h2>사양 비교</h2>
<table border="1">
<tr>{headers_html}</tr>
{rows_html}
</table>
</body>
</html>"""

    path.write_text(html, encoding="utf-8")
    return path


# ---------------------------------------------------------------------------
# T011: create_txt
# ---------------------------------------------------------------------------


def create_txt(path: Path) -> Path:
    """Create a plain .txt file. No external library needed.

    Args:
        path: File path to create.

    Returns:
        The created file Path.
    """
    lines = [WRITER_TITLE, "=" * len(WRITER_TITLE), ""]
    for para in WRITER_PARAGRAPHS:
        lines.append(para)
        lines.append("")
    lines.append("사양 비교")
    lines.append("-" * 10)
    # Simple table
    lines.append("\t".join(WRITER_TABLE_HEADERS))
    for row in WRITER_TABLE_ROWS:
        lines.append("\t".join(row))

    path.write_text("\n".join(lines), encoding="utf-8")
    return path


# ---------------------------------------------------------------------------
# Edge case helpers (Phase 7)
# ---------------------------------------------------------------------------


def create_empty_docx(path: Path) -> Path:
    """Create an empty .docx file with no content (FR-015)."""
    from docx import Document

    doc = Document()
    doc.save(str(path))
    return path


def create_unicode_docx(path: Path) -> Path:
    """Create a .docx with multilingual Unicode text (FR-015).

    Contains Korean, Japanese, Arabic, and emoji characters.
    """
    from docx import Document

    doc = Document()
    doc.add_heading("유니코드 테스트 문서", level=1)
    doc.add_paragraph("한국어: 안녕하세요, LibreFormer 테스트입니다.")
    doc.add_paragraph("日本語: こんにちは、テストです。")
    doc.add_paragraph("العربية: مرحبا، هذا اختبار.")
    doc.add_paragraph("Emoji: 🚀📄✅🎉🔥💻🌍")
    doc.add_paragraph("Mixed: Hello 세계 世界 🌏")

    doc.save(str(path))
    return path


def create_special_chars_txt(path: Path) -> Path:
    """Create a .txt with special characters, tabs, newlines (FR-015)."""
    content = (
        "Special Characters Test\n"
        "=======================\n\n"
        "Tabs:\tFirst\tSecond\tThird\n"
        "Quotes: \"double\" and 'single'\n"
        "Ampersand: A & B\n"
        "Angle brackets: <tag> and </tag>\n"
        "Backslash: C:\\Users\\test\n"
        "Unicode symbols: © ® ™ § ¶ † ‡ • … — –\n"
        "Currency: $ € £ ¥ ₩\n"
        "Math: ± × ÷ ≠ ≤ ≥ ∞\n"
        "Null and control: \x00 \x01 \x02\n"
        "Line endings: CR+LF\r\n"
        "Vertical tab: \x0b\n"
        "Form feed: \x0c\n"
    )
    path.write_bytes(content.encode("utf-8"))
    return path
